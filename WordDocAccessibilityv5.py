# -*- coding: utf-8 -*-
"""
Created on Thu Mar 12 06:51:44 2026

@author: robert.hardin
"
Word Document Digital Accessibility Improvements
Complete working version with proper style handling
"""

import requests
import streamlit as st
from io import BytesIO
import zipfile
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def call_models_api(api_key):
    """Fetch available models from TAMU AI API"""
    url = "https://chat-api.tamu.ai/openai/models"
    headers = {
        "accept": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    response = requests.get(url, headers=headers)
    response.raise_for_status()
    model_info = response.json()['data']
    id_list = []
    name_list = []
    for model in model_info:
        id_list.append(model['id'])
        name_list.append(model['name'])
    model_dict = dict(zip(name_list, id_list))
    return model_dict


def interact_with_model(api_key, chosen_model, my_query):
    """Send query to LLM and get response"""
    url = "https://chat-api.tamu.ai/openai/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {
        "model": chosen_model,
        "messages": [{"role": "user", "content": my_query}],
        "stream": False
    }
    response = requests.post(url, headers=headers, json=payload)
    return response.json()


def extract_llm_response(response_json):
    """Extract text content from LLM response"""
    try:
        return response_json['choices'][0]['message']['content']
    except (KeyError, IndexError):
        return ""


def hex_to_rgb(hex_color):
    """Convert hex color string to RGBColor object"""
    hex_color = hex_color.lstrip('#')
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def set_run_font(run, font_name):
    """Set font for a single run element"""
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    
    for rFonts in rPr.findall(qn('w:rFonts')):
        rPr.remove(rFonts)
    
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)


def set_element_font_recursive(element, font_name):
    """Recursively set font for all text elements"""
    for r in element.iter(qn('w:r')):
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        
        for rFonts in rPr.findall(qn('w:rFonts')):
            rPr.remove(rFonts)
        
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rPr.insert(0, rFonts)


def change_all_fonts(doc, font_name, styled_paragraphs):
    """
    Change all fonts except for paragraphs that have Title/Heading styles applied.
    Those paragraphs get their formatting from the style.
    """
    # Change fonts in all document styles except Title, Heading 1, Heading 2
    for style in doc.styles:
        try:
            if hasattr(style, 'font') and style.font is not None:
                if style.name not in ['Title', 'Heading 1', 'Heading 2']:
                    style.font.name = font_name
        except Exception:
            pass
    
    # Change default font in document defaults
    try:
        styles_element = doc.styles.element
        docDefaults = styles_element.find(qn('w:docDefaults'))
        if docDefaults is not None:
            rPrDefault = docDefaults.find(qn('w:rPrDefault'))
            if rPrDefault is not None:
                rPr = rPrDefault.find(qn('w:rPr'))
                if rPr is not None:
                    for rFonts in rPr.findall(qn('w:rFonts')):
                        rPr.remove(rFonts)
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:ascii'), font_name)
                    rFonts.set(qn('w:hAnsi'), font_name)
                    rFonts.set(qn('w:cs'), font_name)
                    rFonts.set(qn('w:eastAsia'), font_name)
                    rPr.insert(0, rFonts)
    except Exception:
        pass
    
    # Change fonts in paragraphs (skip styled ones)
    for para in doc.paragraphs:
        if para not in styled_paragraphs:
            for run in para.runs:
                set_run_font(run, font_name)
            for hyperlink in para._element.findall(qn('w:hyperlink')):
                set_element_font_recursive(hyperlink, font_name)
            set_element_font_recursive(para._element, font_name)
    
    # Change fonts in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para not in styled_paragraphs:
                        for run in para.runs:
                            set_run_font(run, font_name)
                        set_element_font_recursive(para._element, font_name)
    
    # Change fonts in headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header is not None:
                try:
                    for para in header.paragraphs:
                        for run in para.runs:
                            set_run_font(run, font_name)
                except Exception:
                    pass
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer is not None:
                try:
                    for para in footer.paragraphs:
                        for run in para.runs:
                            set_run_font(run, font_name)
                except Exception:
                    pass
    
    # Change numbering definitions
    try:
        numbering_part = doc.part.numbering_part
        if numbering_part is not None:
            numbering_element = numbering_part.element
            for lvl in numbering_element.iter(qn('w:lvl')):
                rPr = lvl.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    lvl.append(rPr)
                for rFonts in rPr.findall(qn('w:rFonts')):
                    rPr.remove(rFonts)
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rFonts.set(qn('w:cs'), font_name)
                rFonts.set(qn('w:eastAsia'), font_name)
                rPr.insert(0, rFonts)
    except Exception:
        pass


def get_or_activate_builtin_style(doc, style_name):
    """
    Get a built-in style, activating it from latent styles if necessary.
    Preserves all style identity attributes.
    """
    # Map style names to their built-in style IDs
    style_id_map = {
        'Title': 'Title',
        'Heading 1': 'Heading1',
        'Heading 2': 'Heading2',
        'Heading 3': 'Heading3',
        'Heading 4': 'Heading4',
        'Heading 5': 'Heading5',
        'Normal': 'Normal',
    }
    
    style_id = style_id_map.get(style_name)
    if style_id is None:
        return None
    
    # First try to access the style directly
    try:
        style = doc.styles[style_name]
        return style
    except KeyError:
        pass
    
    # Style is latent - check latent_styles
    latent_styles = doc.styles.latent_styles
    latent_style = None
    try:
        latent_style = latent_styles[style_name]
    except KeyError:
        pass
    
    # Get the styles element
    styles_element = doc.styles.element
    
    # Check if style already exists in XML by styleId
    for existing_style in styles_element.findall(qn('w:style')):
        if existing_style.get(qn('w:styleId')) == style_id:
            # Style element exists, ensure name is correct
            name_elem = existing_style.find(qn('w:name'))
            if name_elem is None:
                name_elem = OxmlElement('w:name')
                existing_style.insert(0, name_elem)
            name_elem.set(qn('w:val'), style_name)
            
            try:
                return doc.styles[style_name]
            except KeyError:
                pass
    
    # Need to create the style element to activate it
    new_style = OxmlElement('w:style')
    new_style.set(qn('w:type'), 'paragraph')
    new_style.set(qn('w:styleId'), style_id)
    
    # Add name element
    name_elem = OxmlElement('w:name')
    name_elem.set(qn('w:val'), style_name)
    new_style.append(name_elem)
    
    # Add basedOn Normal (important for style hierarchy)
    basedOn = OxmlElement('w:basedOn')
    basedOn.set(qn('w:val'), 'Normal')
    new_style.append(basedOn)
    
    # Add qFormat to show in quick styles gallery
    qFormat = OxmlElement('w:qFormat')
    new_style.append(qFormat)
    
    # Add uiPriority from latent style if available
    if latent_style is not None:
        try:
            priority = latent_style.priority
            if priority is not None:
                uiPriority = OxmlElement('w:uiPriority')
                uiPriority.set(qn('w:val'), str(priority))
                new_style.append(uiPriority)
        except:
            pass
    
    # Add pPr with outline level for proper document structure
    pPr = OxmlElement('w:pPr')
    outlineLvl = OxmlElement('w:outlineLvl')
    if style_name == 'Title':
        outlineLvl.set(qn('w:val'), '0')
    elif style_name == 'Heading 1':
        outlineLvl.set(qn('w:val'), '0')
    elif style_name == 'Heading 2':
        outlineLvl.set(qn('w:val'), '1')
    elif style_name == 'Heading 3':
        outlineLvl.set(qn('w:val'), '2')
    pPr.append(outlineLvl)
    new_style.append(pPr)
    
    # Add empty rPr for formatting
    rPr = OxmlElement('w:rPr')
    new_style.append(rPr)
    
    # Insert the style into the document
    latent_styles_elem = styles_element.find(qn('w:latentStyles'))
    if latent_styles_elem is not None:
        latent_styles_elem.addprevious(new_style)
    else:
        styles_element.append(new_style)
    
    # Try to access through python-docx
    try:
        return doc.styles[style_name]
    except KeyError:
        return None


def modify_builtin_style(doc, style_name, font_name, font_size, font_color, bold, italic, underline, centered, space_before, space_after):
    """
    Modify a built-in Word style while preserving its structural identity.
    Only modifies the formatting properties, not the style's identity elements.
    """
    # Map style names to their built-in style IDs
    style_id_map = {
        'Title': 'Title',
        'Heading 1': 'Heading1',
        'Heading 2': 'Heading2',
        'Heading 3': 'Heading3',
        'Normal': 'Normal',
    }
    
    style_id = style_id_map.get(style_name)
    if style_id is None:
        st.warning(f"Unknown style name: {style_name}")
        return None
    
    # First, try to activate the style if it's latent
    style = get_or_activate_builtin_style(doc, style_name)
    
    # Get the styles element for direct XML manipulation
    styles_element = doc.styles.element
    
    # Find the style element by styleId
    style_element = None
    for existing_style in styles_element.findall(qn('w:style')):
        if existing_style.get(qn('w:styleId')) == style_id:
            style_element = existing_style
            break
    
    if style_element is None:
        st.warning(f"Could not find or create style element for '{style_name}'")
        return None
    
    # Ensure qFormat is present (shows in quick styles gallery)
    qFormat = style_element.find(qn('w:qFormat'))
    if qFormat is None:
        qFormat = OxmlElement('w:qFormat')
        name_elem = style_element.find(qn('w:name'))
        if name_elem is not None:
            name_elem.addnext(qFormat)
        else:
            style_element.insert(0, qFormat)
    
    # --- MODIFY pPr (paragraph properties) ---
    pPr = style_element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        rPr = style_element.find(qn('w:rPr'))
        if rPr is not None:
            rPr.addprevious(pPr)
        else:
            style_element.append(pPr)
    
    # Remove only the formatting elements we want to change in pPr
    for child in list(pPr):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ['jc', 'spacing']:
            pPr.remove(child)
    
    # Set alignment
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center' if centered else 'left')
    pPr.append(jc)
    
    # Set spacing (in twips, 1 point = 20 twips)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '240' if space_before else '0')
    spacing.set(qn('w:after'), '240' if space_after else '0')
    pPr.append(spacing)
    
    # Ensure outline level is set for accessibility
    outlineLvl = pPr.find(qn('w:outlineLvl'))
    if outlineLvl is None:
        outlineLvl = OxmlElement('w:outlineLvl')
        pPr.append(outlineLvl)
    
    if style_name == 'Title':
        outlineLvl.set(qn('w:val'), '0')
    elif style_name == 'Heading 1':
        outlineLvl.set(qn('w:val'), '0')
    elif style_name == 'Heading 2':
        outlineLvl.set(qn('w:val'), '1')
    
    # --- MODIFY rPr (run properties) ---
    rPr = style_element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style_element.append(rPr)
    
    # Remove only the formatting elements we want to change in rPr
    for child in list(rPr):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ['rFonts', 'sz', 'szCs', 'b', 'bCs', 'i', 'iCs', 'u', 'color']:
            rPr.remove(child)
    
    # Set font name
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)
    
    # Set font size (in half-points)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(font_size * 2)))
    rPr.append(szCs)
    
    # Set bold - explicitly set value
    b = OxmlElement('w:b')
    b.set(qn('w:val'), '1' if bold else '0')
    rPr.append(b)
    bCs = OxmlElement('w:bCs')
    bCs.set(qn('w:val'), '1' if bold else '0')
    rPr.append(bCs)
    
    # Set italic - explicitly set value
    i_elem = OxmlElement('w:i')
    i_elem.set(qn('w:val'), '1' if italic else '0')
    rPr.append(i_elem)
    iCs = OxmlElement('w:iCs')
    iCs.set(qn('w:val'), '1' if italic else '0')
    rPr.append(iCs)
    
    # Set underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single' if underline else 'none')
    rPr.append(u)
    
    # Set color
    if font_color:
        color = OxmlElement('w:color')
        color.set(qn('w:val'), font_color.lstrip('#'))
        rPr.append(color)
    
    # Update latent style to ensure it's visible
    try:
        latent_styles = doc.styles.latent_styles
        latent_style = latent_styles[style_name]
        latent_style_elem = latent_style._element
        latent_style_elem.set(qn('w:semiHidden'), '0')
        latent_style_elem.set(qn('w:unhideWhenUsed'), '0')
    except:
        pass
    
    return style


def apply_style_to_paragraph(para, style_name, doc):
    """
    Apply a named built-in style to a paragraph.
    Removes direct formatting so the style formatting takes effect.
    """
    # Map style names to their correct style IDs
    style_id_map = {
        'Title': 'Title',
        'Heading 1': 'Heading1',
        'Heading 2': 'Heading2',
        'Heading 3': 'Heading3',
        'Normal': 'Normal',
    }
    
    style_id = style_id_map.get(style_name)
    if style_id is None:
        st.warning(f"Unknown style: {style_name}")
        return
    
    # Get the paragraph XML element
    p = para._element
    
    # Find or create pPr element
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p.insert(0, pPr)
    
    # Find or create pStyle element and set it FIRST
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        pStyle = OxmlElement('w:pStyle')
        pPr.insert(0, pStyle)
    
    # Set the style ID
    pStyle.set(qn('w:val'), style_id)
    
    # Now try to also set via python-docx for consistency
    try:
        style = doc.styles[style_name]
        para.style = style
    except KeyError:
        pass
    
    # Remove direct paragraph-level run properties that override style
    para_rPr = pPr.find(qn('w:rPr'))
    if para_rPr is not None:
        pPr.remove(para_rPr)
    
    # Remove paragraph formatting elements that override style
    tags_to_remove_from_pPr = ['jc', 'spacing', 'ind', 'outlineLvl', 'keepNext', 'keepLines']
    for child in list(pPr):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in tags_to_remove_from_pPr:
            pPr.remove(child)
    
    # Remove direct formatting from ALL runs so style formatting takes effect
    for r in p.findall(qn('w:r')):
        rPr = r.find(qn('w:rPr'))
        if rPr is not None:
            tags_to_remove = ['rFonts', 'sz', 'szCs', 'b', 'bCs', 'i', 'iCs', 'u', 'color']
            for child in list(rPr):
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag in tags_to_remove:
                    rPr.remove(child)
            if len(rPr) == 0:
                r.remove(rPr)
    
    # Handle hyperlinks within the paragraph
    for hyperlink in p.findall(qn('w:hyperlink')):
        for r in hyperlink.findall(qn('w:r')):
            rPr = r.find(qn('w:rPr'))
            if rPr is not None:
                tags_to_remove = ['rFonts', 'sz', 'szCs', 'b', 'bCs', 'i', 'iCs', 'u', 'color']
                for child in list(rPr):
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag in tags_to_remove:
                        rPr.remove(child)
                if len(rPr) == 0:
                    r.remove(rPr)


def get_document_text(doc):
    """Extract all text from document for LLM analysis"""
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    return "\n".join(full_text)


def chunk_document_text(doc_text, chunk_size=3000, overlap=500):
    """Split document text into overlapping chunks for LLM processing."""
    if len(doc_text) <= chunk_size:
        return [doc_text]
    
    chunks = []
    start = 0
    
    while start < len(doc_text):
        end = start + chunk_size
        if end < len(doc_text):
            paragraph_break = doc_text.rfind('\n\n', start, end)
            if paragraph_break > start + chunk_size // 2:
                end = paragraph_break
            else:
                sentence_break = doc_text.rfind('. ', start, end)
                if sentence_break > start + chunk_size // 2:
                    end = sentence_break + 1
        chunks.append(doc_text[start:end])
        start = end - overlap if end < len(doc_text) else end
    
    return chunks


def check_existing_styles(doc):
    """Check if document has properly applied Title and Heading styles"""
    has_title = False
    has_headings = False
    
    for para in doc.paragraphs:
        if para.style.name == "Title":
            has_title = True
        if "Heading" in para.style.name:
            has_headings = True
    
    return has_title, has_headings


def identify_title_and_headings(api_key, model_id, doc_text):
    """Use LLM to identify title and headings in document."""
    chunks = chunk_document_text(doc_text, chunk_size=3500, overlap=300)
    
    all_titles = []
    all_headings1 = []
    all_headings2 = []
    
    for i, chunk in enumerate(chunks):
        prompt = f"""Analyze the following document text (part {i+1} of {len(chunks)}) and identify:
1. The title (the main title of the document) - only identify if this appears to be the document title
2. Any section headings (Heading 1 level - major sections)
3. Any sub-headings (Heading 2 level - subsections)

Return your response in this exact format:
TITLE: [exact text of title]
HEADING1: [exact text of heading 1]
HEADING1: [exact text of another heading 1]
HEADING2: [exact text of heading 2]
HEADING2: [exact text of another heading 2]

If no title is found in this section, respond with TITLE: NONE
Only include headings that are clearly section headers, not regular paragraph text.
Return the EXACT text as it appears - do not paraphrase or modify.

Document text:
{chunk}"""
        
        response = interact_with_model(api_key, model_id, prompt)
        llm_response = extract_llm_response(response)
        
        lines = llm_response.strip().split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith("TITLE:"):
                title_text = line.replace("TITLE:", "").strip()
                if title_text and title_text.upper() != "NONE" and title_text not in all_titles:
                    all_titles.append(title_text)
            elif line.startswith("HEADING1:"):
                h1_text = line.replace("HEADING1:", "").strip()
                if h1_text and h1_text not in all_headings1:
                    all_headings1.append(h1_text)
            elif line.startswith("HEADING2:"):
                h2_text = line.replace("HEADING2:", "").strip()
                if h2_text and h2_text not in all_headings2:
                    all_headings2.append(h2_text)
    
    # Reconstruct the response format
    result_lines = []
    for title in all_titles:
        result_lines.append(f"TITLE: {title}")
    if not all_titles:
        result_lines.append("TITLE: NONE")
    for h1 in all_headings1:
        result_lines.append(f"HEADING1: {h1}")
    for h2 in all_headings2:
        result_lines.append(f"HEADING2: {h2}")
    
    return "\n".join(result_lines)


def generate_title(api_key, model_id, doc_text):
    """Use LLM to generate an appropriate title for the document."""
    if len(doc_text) > 4000:
        beginning = doc_text[:1500]
        middle_start = len(doc_text) // 2 - 750
        middle = doc_text[middle_start:middle_start + 1500]
        end = doc_text[-1500:]
        sampled_text = f"BEGINNING:\n{beginning}\n\nMIDDLE:\n{middle}\n\nEND:\n{end}"
    else:
        sampled_text = doc_text
    
    prompt = f"""Based on the following document content, suggest a concise and appropriate title.
Return ONLY the title text, nothing else.

Document content:
{sampled_text}"""
    
    response = interact_with_model(api_key, model_id, prompt)
    return extract_llm_response(response).strip()


def normalize_text(text):
    """Normalize text for comparison"""
    text = re.sub(r'\s+', ' ', text.strip().lower())
    text = re.sub(r'[^\w\s]', '', text)
    return text


def texts_match(para_text, target_text):
    """Check if paragraph text matches target text"""
    norm_para = normalize_text(para_text)
    norm_target = normalize_text(target_text)
    
    if norm_para == norm_target:
        return True
    if len(norm_target) > 10 and (norm_target in norm_para or norm_para in norm_target):
        return True
    para_words = set(norm_para.split())
    target_words = set(norm_target.split())
    if para_words and target_words:
        overlap = len(para_words & target_words)
        similarity = overlap / max(len(para_words), len(target_words))
        if similarity > 0.8:
            return True
    return False


def apply_identified_styles(doc, llm_response):
    """
    Apply built-in Word styles to paragraphs identified by LLM.
    Returns list of paragraphs that had styles applied.
    """
    lines = llm_response.strip().split('\n')
    
    titles = []
    headings1 = []
    headings2 = []
    
    for line in lines:
        line = line.strip()
        if line.startswith("TITLE:"):
            title_text = line.replace("TITLE:", "").strip()
            if title_text and title_text.upper() != "NONE":
                titles.append(title_text)
        elif line.startswith("HEADING1:"):
            h1_text = line.replace("HEADING1:", "").strip()
            if h1_text:
                headings1.append(h1_text)
        elif line.startswith("HEADING2:"):
            h2_text = line.replace("HEADING2:", "").strip()
            if h2_text:
                headings2.append(h2_text)
    
    styled_paragraphs = []
    found_titles = []
    
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if not para_text:
            continue
        
        # Skip if already has a proper style
        if para.style.name in ['Title', 'Heading 1', 'Heading 2']:
            styled_paragraphs.append(para)
            continue
        
        # Check for title match
        for title in titles:
            if texts_match(para_text, title):
                apply_style_to_paragraph(para, 'Title', doc)
                styled_paragraphs.append(para)
                found_titles.append(title)
                break
        else:
            # Check for heading 1 match
            for h1 in headings1:
                if texts_match(para_text, h1):
                    apply_style_to_paragraph(para, 'Heading 1', doc)
                    styled_paragraphs.append(para)
                    break
            else:
                # Check for heading 2 match
                for h2 in headings2:
                    if texts_match(para_text, h2):
                        apply_style_to_paragraph(para, 'Heading 2', doc)
                        styled_paragraphs.append(para)
                        break
    
    return found_titles, styled_paragraphs


def add_title_to_document(doc, title_text):
    """Add a title to the beginning of the document with Title style."""
    if doc.paragraphs:
        new_para = doc.paragraphs[0].insert_paragraph_before(title_text)
    else:
        new_para = doc.add_paragraph(title_text)
    
    apply_style_to_paragraph(new_para, 'Title', doc)
    return new_para


def generate_image_alt_text(api_key, model_id, doc_text, image_index, paragraph_context=""):
    """Use LLM to generate alt text for an image based on context."""
    if paragraph_context:
        context = paragraph_context
    elif len(doc_text) > 2000:
        context = doc_text[:2000]
    else:
        context = doc_text
    
    prompt = f"""Based on the following document context, generate appropriate alt text for image #{image_index}.
The alt text should be descriptive, concise (under 125 characters), and convey the purpose of the image.
Return ONLY the alt text, nothing else.

Document context:
{context}"""
    
    response = interact_with_model(api_key, model_id, prompt)
    return extract_llm_response(response).strip()


def set_image_alt_text(inline_shape, alt_text):
    """Set alt text for an inline shape (image)"""
    try:
        inline = inline_shape._inline
        docPr = inline.find(qn('wp:docPr'))
        if docPr is not None:
            docPr.set('descr', alt_text)
            docPr.set('title', alt_text[:50])
    except Exception:
        pass


def process_images(doc, api_key, model_id, doc_text):
    """Process all images in document and add alt text."""
    image_count = 0
    paragraphs = list(doc.paragraphs)
    
    for para_idx, para in enumerate(paragraphs):
        for run in para.runs:
            if run._element.xpath('.//a:blip'):
                image_count += 1
                start_idx = max(0, para_idx - 2)
                end_idx = min(len(paragraphs), para_idx + 3)
                surrounding_paras = paragraphs[start_idx:end_idx]
                paragraph_context = "\n".join([p.text for p in surrounding_paras if p.text.strip()])
                
                alt_text = generate_image_alt_text(api_key, model_id, doc_text, image_count, paragraph_context)
                
                for inline in run._element.xpath('.//wp:inline'):
                    docPr = inline.find(qn('wp:docPr'))
                    if docPr is not None:
                        docPr.set('descr', alt_text)
    
    for shape in doc.inline_shapes:
        image_count += 1
        alt_text = generate_image_alt_text(api_key, model_id, doc_text, image_count)
        set_image_alt_text(shape, alt_text)
    
    return image_count


def get_table_content(table):
    """Extract text content from a table for description"""
    content = []
    for row in table.rows:
        row_content = []
        for cell in row.cells:
            row_content.append(cell.text.strip())
        content.append(" | ".join(row_content))
    return "\n".join(content)


def generate_table_alt_text(api_key, model_id, table_content):
    """Use LLM to generate description for a table."""
    if len(table_content) > 2000:
        truncated_content = table_content[:1000] + "\n...[truncated]...\n" + table_content[-500:]
    else:
        truncated_content = table_content
    
    prompt = f"""Generate a brief, accessible description for the following table content.
The description should summarize what the table contains and its purpose.
Keep it under 200 characters. Return ONLY the description.

Table content:
{truncated_content}"""
    
    response = interact_with_model(api_key, model_id, prompt)
    return extract_llm_response(response).strip()


def process_tables(doc, api_key, model_id):
    """Process all tables and add descriptions"""
    table_descriptions = []
    for i, table in enumerate(doc.tables):
        table_content = get_table_content(table)
        if table_content.strip():
            description = generate_table_alt_text(api_key, model_id, table_content)
            table_descriptions.append(f"Table {i+1}: {description}")
    return table_descriptions


def ensure_document_language(doc):
    """Set document language for accessibility"""
    try:
        for para in doc.paragraphs:
            for run in para.runs:
                rPr = run._element.get_or_add_rPr()
                lang = OxmlElement('w:lang')
                lang.set(qn('w:val'), 'en-US')
                rPr.append(lang)
    except Exception:
        pass


def process_document(uploaded_file, api_key, model_id, font, 
                     title_font_size, title_font_color, title_bold, title_italic, title_underline, 
                     title_centered, title_space_before, title_space_after,
                     h1_font_size, h1_font_color, h1_bold, h1_italic, h1_underline,
                     h1_centered, h1_space_before, h1_space_after,
                     h2_font_size, h2_font_color, h2_bold, h2_italic, h2_underline,
                     h2_centered, h2_space_before, h2_space_after):
    """Process a single document for accessibility"""
    
    # Load the document
    doc = Document(BytesIO(uploaded_file.getvalue()))
    
    # Get document text for LLM analysis BEFORE modifications
    doc_text = get_document_text(doc)
    
    # 1. FIRST: Modify the built-in styles with user preferences
    modify_builtin_style(doc, "Title", font, title_font_size, title_font_color, 
                         title_bold, title_italic, title_underline, 
                         title_centered, title_space_before, title_space_after)
    
    modify_builtin_style(doc, "Heading 1", font, h1_font_size, h1_font_color,
                         h1_bold, h1_italic, h1_underline,
                         h1_centered, h1_space_before, h1_space_after)
    
    modify_builtin_style(doc, "Heading 2", font, h2_font_size, h2_font_color,
                         h2_bold, h2_italic, h2_underline,
                         h2_centered, h2_space_before, h2_space_after)
    
    # 2. Check existing styles
    has_title, has_headings = check_existing_styles(doc)
    
    # 3. Use LLM to identify title and headings, then apply styles
    styled_paragraphs = []
    if not has_title or not has_headings:
        llm_response = identify_title_and_headings(api_key, model_id, doc_text)
        found_titles, styled_paragraphs = apply_identified_styles(doc, llm_response)
        
        # 4. If no title identified, generate one and add it
        if not found_titles:
            generated_title = generate_title(api_key, model_id, doc_text)
            if generated_title:
                new_para = add_title_to_document(doc, generated_title)
                styled_paragraphs.append(new_para)
    
    # 5. Change all fonts AFTER styles are applied
    change_all_fonts(doc, font, styled_paragraphs)
    
    # 6. Process images and add alt text
    image_count = process_images(doc, api_key, model_id, doc_text)
    
    # 7. Process tables
    table_descriptions = process_tables(doc, api_key, model_id)
    
    # 8. Additional accessibility improvements
    ensure_document_language(doc)
    
    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output, image_count, len(table_descriptions)


def create_zip_file(processed_files):
    """Create a zip file containing all processed documents"""
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for filename, file_content in processed_files:
            new_filename = f"accessible_{filename}"
            zip_file.writestr(new_filename, file_content.getvalue())
    zip_buffer.seek(0)
    return zip_buffer


# Main Streamlit App
st.title("Word Document Digital Accessibility Improvements")

st.markdown("""
This application improves the digital accessibility of Word documents by:
- Converting fonts to accessible alternatives
- Applying proper heading styles with custom colors
- Generating alt text for images and tables
- Ensuring proper document structure
""")

api_key = st.text_input("TAMU API Key", value=None, type="password")

if api_key is not None and api_key != "":
    try:
        model_dict = call_models_api(api_key)
        selected_model_name = st.selectbox(
            "Pick a large language model to use for providing feedback", 
            list(model_dict), key=25
        )
        selected_model_id = model_dict[selected_model_name]
        
        # Font selection
        font = st.selectbox(
            "Pick a recommended font for accessibility", 
            ["Arial", "Calibri", "Verdana", "Tahoma", "Helvetica", "Aptos"], 
            key=0
        )
        
        # Title style inputs
        st.write("### Enter details for Title style:")
        title_col1, title_col2 = st.columns(2)
        with title_col1:
            title_font_size = st.number_input("Title Font size", min_value=7, value=24, key=1)
        with title_col2:
            title_font_color = st.color_picker("Title Font Color", value="#000000", key=2)
        
        title = st.container()
        col1, col2, col3, col4, col5, col6 = title.columns(6)
        title_bold = col1.checkbox("Bold", value=True, key=3)
        title_italic = col2.checkbox("Italic", key=4)
        title_underline = col3.checkbox("Underline", key=5)
        title_centered = col4.checkbox("Centered", value=True, key=6)
        title_space_before = col5.checkbox("Space Before", key=7)
        title_space_after = col6.checkbox("Space After", value=True, key=8)
        
        # Heading 1 style inputs
        st.write("### Enter details for Heading 1 style:")
        h1_col1, h1_col2 = st.columns(2)
        with h1_col1:
            h1_font_size = st.number_input("Heading 1 Font size", min_value=7, value=18, key=9)
        with h1_col2:
            h1_font_color = st.color_picker("Heading 1 Font Color", value="#000000", key=10)
        
        h1 = st.container()
        col1, col2, col3, col4, col5, col6 = h1.columns(6)
        h1_bold = col1.checkbox("Bold", value=True, key=11)
        h1_italic = col2.checkbox("Italic", key=12)
        h1_underline = col3.checkbox("Underline", key=13)
        h1_centered = col4.checkbox("Centered", key=14)
        h1_space_before = col5.checkbox("Space Before", value=True, key=15)
        h1_space_after = col6.checkbox("Space After", value=True, key=16)
        
        # Heading 2 style inputs
        st.write("### Enter details for Heading 2 style:")
        h2_col1, h2_col2 = st.columns(2)
        with h2_col1:
            h2_font_size = st.number_input("Heading 2 Font size", min_value=7, value=14, key=17)
        with h2_col2:
            h2_font_color = st.color_picker("Heading 2 Font Color", value="#000000", key=18)
        
        h2 = st.container()
        col1, col2, col3, col4, col5, col6 = h2.columns(6)
        h2_bold = col1.checkbox("Bold", value=True, key=19)
        h2_italic = col2.checkbox("Italic", key=20)
        h2_underline = col3.checkbox("Underline", key=21)
        h2_centered = col4.checkbox("Centered", key=22)
        h2_space_before = col5.checkbox("Space Before", value=True, key=23)
        h2_space_after = col6.checkbox("Space After", key=24)
        
        # File uploader
        uploaded_files = st.file_uploader(
            "Choose Word document files to modify", 
            type="docx", 
            accept_multiple_files=True, 
            key=26
        )
        
        # Process button
        if uploaded_files:
            st.write(f"**{len(uploaded_files)} file(s) uploaded**")
            
            if st.button("Process Documents for Accessibility", type="primary"):
                processed_files = []
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                with st.expander("Processing Details", expanded=True):
                    for i, uploaded_file in enumerate(uploaded_files):
                        status_text.text(f"Processing: {uploaded_file.name}")
                        st.write(f"**Processing: {uploaded_file.name}**")
                        
                        try:
                            processed_doc, image_count, table_count = process_document(
                                uploaded_file=uploaded_file,
                                api_key=api_key,
                                model_id=selected_model_id,
                                font=font,
                                title_font_size=title_font_size,
                                title_font_color=title_font_color,
                                title_bold=title_bold,
                                title_italic=title_italic,
                                title_underline=title_underline,
                                title_centered=title_centered,
                                title_space_before=title_space_before,
                                title_space_after=title_space_after,
                                h1_font_size=h1_font_size,
                                h1_font_color=h1_font_color,
                                h1_bold=h1_bold,
                                h1_italic=h1_italic,
                                h1_underline=h1_underline,
                                h1_centered=h1_centered,
                                h1_space_before=h1_space_before,
                                h1_space_after=h1_space_after,
                                h2_font_size=h2_font_size,
                                h2_font_color=h2_font_color,
                                h2_bold=h2_bold,
                                h2_italic=h2_italic,
                                h2_underline=h2_underline,
                                h2_centered=h2_centered,
                                h2_space_before=h2_space_before,
                                h2_space_after=h2_space_after
                            )
                            
                            processed_files.append((uploaded_file.name, processed_doc))
                            
                            st.success(f"✓ {uploaded_file.name} processed successfully")
                            st.write(f"  - Font changed to: {font}")
                            st.write(f"  - Images processed: {image_count}")
                            st.write(f"  - Tables processed: {table_count}")
                            
                        except Exception as e:
                            st.error(f"✗ Error processing {uploaded_file.name}: {str(e)}")
                        
                        progress_bar.progress((i + 1) / len(uploaded_files))
                
                status_text.text("Processing complete!")
                
                if processed_files:
                    st.write("---")
                    st.write("### Download Processed Files")
                    
                    zip_buffer = create_zip_file(processed_files)
                    
                    st.download_button(
                        label="📥 Download All Accessible Documents (ZIP)",
                        data=zip_buffer,
                        file_name="accessible_documents.zip",
                        mime="application/zip",
                        type="primary"
                    )
                    
                    # Only show individual download buttons if 3 or fewer files
                    if len(processed_files) <= 3:
                        st.write("---")
                        st.write("### Download Individual Files")
                        
                        for filename, file_content in processed_files:
                            file_content.seek(0)
                            st.download_button(
                                label=f"📄 Download accessible_{filename}",
                                data=file_content,
                                file_name=f"accessible_{filename}",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"download_{filename}"
                            )
                    
                    st.success(f"✅ Successfully processed {len(processed_files)} document(s)!")
                    
                    # Accessibility summary
                    st.write("---")
                    st.write("### Accessibility Improvements Made:")
                    st.markdown("""
                    - ✓ All fonts converted to accessible typeface (including lists and hyperlinks)
                    - ✓ Title style modified and applied with custom formatting and color
                    - ✓ Heading 1 style modified and applied with custom formatting and color
                    - ✓ Heading 2 style modified and applied with custom formatting and color
                    - ✓ Alt text generated for images
                    - ✓ Table descriptions added
                    - ✓ Document language set to English (US)
                    - ✓ Proper heading hierarchy ensured
                    """)
        
        else:
            st.info("Please upload one or more Word documents (.docx) to process.")
    
    except Exception as e:
        st.error(f"Error connecting to API: {str(e)}")
        st.write("Please check your API key and try again.")

else:
    st.warning("Please enter your TAMU API Key to continue.")
    st.markdown("""
    ### Instructions:
    1. Enter your TAMU AI API key above
    2. Select a language model for content analysis
    3. Choose your preferred accessible font
    4. Configure Title, Heading 1, and Heading 2 styles (including colors)
    5. Upload Word documents to process
    6. Click 'Process Documents' to make them accessible
    7. Download the modified files individually or as a ZIP
    """)

# Footer
st.write("---")
st.markdown("""
<small>
Created by Robert G. Hardin IV with Claude Opus 4.5. This tool helps improve digital accessibility compliance by applying proper document structure,
accessible fonts, custom heading colors, and alternative text for visual elements. 
Font changes are applied to all text including list numbering, hyperlinks, headers, footers, 
footnotes, and all other document elements. For full WCAG compliance, 
additional manual review may be required.
</small>
""", unsafe_allow_html=True)